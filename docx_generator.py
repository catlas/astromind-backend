#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOCX Generator for Astrology Reports
Generates professional DOCX documents for periods > 6 months
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from datetime import datetime
import re


class DOCXGenerator:
    def __init__(self):
        self.planet_names = {
            'Sun': 'Слънце', 'Moon': 'Луна', 'Mercury': 'Меркурий',
            'Venus': 'Венера', 'Mars': 'Марс', 'Jupiter': 'Юпитер',
            'Saturn': 'Сатурн', 'Uranus': 'Уран', 'Neptune': 'Нептун',
            'Pluto': 'Плутон', 'Node': 'Възходящ Възел', 'Chiron': 'Хирон'
        }
        
        self.sign_names = {
            'Aries': 'Овен', 'Taurus': 'Телец', 'Gemini': 'Близнаци',
            'Cancer': 'Рак', 'Leo': 'Лъв', 'Virgo': 'Дева',
            'Libra': 'Везни', 'Scorpio': 'Скорпион', 'Sagittarius': 'Стрелец',
            'Capricorn': 'Козирог', 'Aquarius': 'Водолей', 'Pisces': 'Риби'
        }
        
        self.aspect_names = {
            'conjunction': 'съвпад', 'sextile': 'секстил',
            'square': 'квадратура', 'trine': 'тригон', 'opposition': 'опозиция'
        }
    
    def generate_docx(self, data: dict) -> bytes:
        """
        Generate DOCX report from chart data
        
        Args:
            data: Dictionary containing:
                - user_name: str
                - birth_date: str
                - birth_time: str
                - birth_city: str
                - report_type: str
                - natal_chart: dict
                - natal_aspects: list
                - monthly_results: list[dict{month: str, text: str}]
        
        Returns:
            bytes: DOCX file content
        """
        doc = Document()
        
        # Set default font and spacing (for content pages)
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(9)  # Reduced from 13 to 9
        
        # Set paragraph spacing - reduced
        paragraph_format = style.paragraph_format
        paragraph_format.space_after = Pt(4)  # Reduced from 12
        paragraph_format.line_spacing = 1.3  # Reduced from 1.8
        
        # 1. Cover Page
        self._add_cover_page(doc, data)
        doc.add_page_break()
        
        # 2. Chart Summary
        if data.get('natal_chart'):
            self._add_chart_summary(doc, data)
            doc.add_page_break()
        
        # 3. Monthly Interpretations
        monthly_results = data.get('monthly_results', [])
        for idx, month_data in enumerate(monthly_results):
            self._add_month_section(doc, month_data)
            # Only add page break if NOT the last month
            if idx < len(monthly_results) - 1:
                doc.add_page_break()
        
        # Add footer to all sections
        self._add_footer(doc, data)
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def _add_cover_page(self, doc, data):
        """Create professional cover page"""
        # Main title
        title = doc.add_heading('АСТРОЛОГИЧЕН ДОКЛАД', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        
        # Subtitle - ensure UTF-8 string (no spacer)
        user_name = str(data.get("user_name", "Неизвестен"))
        subtitle = doc.add_heading(f'Подготвен за: {user_name}', level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(18)
        
        # Birth info - ensure UTF-8 strings (no spacer)
        birth_para = doc.add_paragraph()
        birth_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        birth_date = str(data.get("birth_date", ""))
        birth_time = str(data.get("birth_time", ""))
        birth_city = str(data.get("birth_city", ""))
        birth_para.add_run(f'Дата на раждане: {birth_date} в {birth_time}\n')
        birth_para.add_run(f'Място: {birth_city}')
        
        # Separator (no spacers)
        doc.add_paragraph('_' * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Report type - ensure UTF-8 string (no spacer)
        type_para = doc.add_paragraph()
        type_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        report_type = str(data.get("report_type", "Астрологичен Анализ"))
        type_run = type_para.add_run(f'Тип анализ: {report_type}')
        type_run.font.size = Pt(14)
        type_run.font.color.rgb = RGBColor(100, 100, 100)
    
    def _add_chart_summary(self, doc, data):
        """Add chart summary section"""
        natal_chart = data.get('natal_chart', {})
        natal_aspects = data.get('natal_aspects', [])
        
        # Section title - reduced size
        heading = doc.add_heading('Обобщена информация за картата', level=2)
        heading_run = heading.runs[0]
        heading_run.font.size = Pt(14)  # Reduced from 20 to 14
        heading.paragraph_format.space_before = Pt(4)
        heading.paragraph_format.space_after = Pt(2)
        
        # 1. Planetary positions
        planets_heading = doc.add_heading('1. ПЛАНЕТАРНИ ПОЗИЦИИ', level=3)
        planets_heading.runs[0].font.size = Pt(11)  # Increased from 10 to 11 (+10%)
        planets_heading.runs[0].font.color.rgb = RGBColor(100, 150, 200)  # Blue color instead of black
        planets_heading.paragraph_format.space_before = Pt(2)
        planets_heading.paragraph_format.space_after = Pt(2)
        
        # Create two-column layout using table
        planets = natal_chart.get('planets', {})
        if planets:
            table = doc.add_table(rows=(len(planets) + 2) // 2, cols=2)
            table.style = 'Table Grid'  # White background for all rows
            
            planet_items = list(planets.items())
            for idx, (planet_name, planet_data) in enumerate(planet_items):
                row_idx = idx // 2
                col_idx = idx % 2
                cell = table.rows[row_idx].cells[col_idx]
                
                planet_bg = self.planet_names.get(planet_name, planet_name)
                position = self._translate_sign(planet_data.get('formatted_pos', ''))
                
                # Add text without bold for planet name
                para = cell.paragraphs[0]
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.space_before = Pt(0)
                
                name_run = para.add_run(f'{planet_bg}: ')
                name_run.font.size = Pt(9)  # Increased from 8 to 9 (+10%)
                name_run.font.bold = False  # NOT bold
                
                pos_run = para.add_run(position)
                pos_run.font.size = Pt(9)  # Increased from 8 to 9 (+10%)
            
            # Add Ascendant if available
            angles = natal_chart.get('angles', {})
            if angles and angles.get('Ascendant') is not None:
                asc_formatted = angles.get('Ascendant_formatted', f"{int(angles.get('Ascendant', 0))}°")
                asc_translated = self._translate_sign(asc_formatted)
                
                # Add to last cell
                last_row = len(planet_items) // 2
                last_col = len(planet_items) % 2
                if last_col == 0:
                    table.add_row()
                    last_row += 1
                    last_col = 0
                
                cell = table.rows[last_row].cells[last_col]
                para = cell.paragraphs[0]
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.space_before = Pt(0)
                
                name_run = para.add_run('Асцендент: ')
                name_run.font.size = Pt(9)  # Increased from 8 to 9 (+10%)
                name_run.font.bold = False
                
                pos_run = para.add_run(asc_translated)
                pos_run.font.size = Pt(9)  # Increased from 8 to 9 (+10%)
        
        # 2. Houses (no spacer)
        houses_heading = doc.add_heading('2. ДОМОВЕ', level=3)
        houses_heading.runs[0].font.size = Pt(11)  # Increased from 10 to 11 (+10%)
        houses_heading.runs[0].font.color.rgb = RGBColor(100, 150, 200)  # Blue color instead of black
        houses_heading.paragraph_format.space_before = Pt(2)
        houses_heading.paragraph_format.space_after = Pt(2)
        
        # Group planets by house
        planets_by_house = {}
        for planet_name, planet_data in planets.items():
            if planet_data and planet_data.get('longitude') is not None:
                house_num = planet_data.get('house', 1)
                if house_num not in planets_by_house:
                    planets_by_house[house_num] = []
                planets_by_house[house_num].append(self.planet_names.get(planet_name, planet_name))
        
        # Create two-column table for houses
        if planets_by_house:
            table = doc.add_table(rows=(12 + 1) // 2, cols=2)
            table.style = 'Table Grid'  # White background for all rows
            
            for house_num in range(1, 13):
                row_idx = (house_num - 1) // 2
                col_idx = (house_num - 1) % 2
                cell = table.rows[row_idx].cells[col_idx]
                
                planets_list = planets_by_house.get(house_num, [])
                if planets_list:
                    suffix = self._get_house_suffix(house_num)
                    planets_str = ', '.join(planets_list)
                    
                    # Add text without bold for house number
                    para = cell.paragraphs[0]
                    para.paragraph_format.space_after = Pt(0)
                    para.paragraph_format.space_before = Pt(0)
                    
                    house_run = para.add_run(f'{house_num}-{suffix} дом: ')
                    house_run.font.size = Pt(9)  # Increased from 8 to 9 (+10%)
                    house_run.font.bold = False  # NOT bold
                    
                    planets_run = para.add_run(planets_str)
                    planets_run.font.size = Pt(9)  # Increased from 8 to 9 (+10%)
                else:
                    # Empty house - show as "празен"
                    suffix = self._get_house_suffix(house_num)
                    para = cell.paragraphs[0]
                    para.paragraph_format.space_after = Pt(0)
                    para.paragraph_format.space_before = Pt(0)
                    
                    house_run = para.add_run(f'{house_num}-{suffix} дом: ')
                    house_run.font.size = Pt(9)  # Increased from 8 to 9 (+10%)
                    house_run.font.bold = False
                    
                    empty_run = para.add_run('празен')
                    empty_run.font.size = Pt(9)  # Increased from 8 to 9 (+10%)
                    empty_run.font.italic = True
                    empty_run.font.color.rgb = RGBColor(150, 150, 150)
        
        # 3. Aspects (no spacer)
        if natal_aspects:
            aspects_heading = doc.add_heading('3. АСПЕКТИ', level=3)
            aspects_heading.runs[0].font.size = Pt(11)  # Increased from 10 to 11 (+10%)
            aspects_heading.runs[0].font.color.rgb = RGBColor(100, 150, 200)  # Blue color instead of black
            aspects_heading.paragraph_format.space_before = Pt(2)
            aspects_heading.paragraph_format.space_after = Pt(2)
            
            # Create three-column table for aspects
            rows_needed = (len(natal_aspects) + 2) // 3
            table = doc.add_table(rows=rows_needed, cols=3)
            table.style = 'Table Grid'  # White background for all rows
            
            for idx, aspect in enumerate(natal_aspects):
                row_idx = idx // 3
                col_idx = idx % 3
                cell = table.rows[row_idx].cells[col_idx]
                
                planet1 = self.planet_names.get(aspect.get('planet1', ''), aspect.get('planet1', ''))
                planet2 = self.planet_names.get(aspect.get('planet2', ''), aspect.get('planet2', ''))
                aspect_name = self.aspect_names.get(aspect.get('aspect', ''), aspect.get('aspect', ''))
                
                # Add text without bold for first planet
                para = cell.paragraphs[0]
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.space_before = Pt(0)
                
                planet1_run = para.add_run(f'{planet1} – {planet2} ')
                planet1_run.font.size = Pt(9)  # Increased from 8 to 9 (+10%)
                planet1_run.font.bold = False  # NOT bold
                
                aspect_run = para.add_run(aspect_name)
                aspect_run.font.size = Pt(9)  # Increased from 8 to 9 (+10%)
    
    def _add_month_section(self, doc, month_data):
        """Add monthly interpretation section"""
        # Ensure UTF-8 strings
        month_name = str(month_data.get('month', 'Месец'))
        month_text = str(month_data.get('text', ''))
        
        # Month title (without emoji for encoding safety)
        # Skip adding month title if it's "Анализ" (static mode - text already has title)
        if month_name != 'Анализ':
            heading = doc.add_heading(month_name, level=1)
            heading_run = heading.runs[0]
            heading_run.font.size = Pt(14)  # Reduced from 22 to 14
            heading_run.font.color.rgb = RGBColor(139, 92, 246)  # Purple
            heading.paragraph_format.space_before = Pt(6)
            heading.paragraph_format.space_after = Pt(4)
        
        # Process and add text with formatting
        self._add_formatted_text(doc, month_text)
    
    def _add_formatted_text(self, doc, text):
        """Add text with markdown-style formatting converted to DOCX"""
        lines = text.split('\n')
        current_paragraph = None
        
        for line in lines:
            line = line.strip()
            
            # Skip empty lines (no extra spacing)
            if not line:
                current_paragraph = None
                continue
            
            # H2 headers (##)
            if line.startswith('## '):
                heading_text = line[3:].strip()
                heading = doc.add_heading(heading_text, level=2)
                heading_run = heading.runs[0]
                heading_run.font.size = Pt(11)  # Reduced from 19 to 11
                heading_run.font.color.rgb = RGBColor(139, 92, 246)
                heading.paragraph_format.space_before = Pt(4)
                heading.paragraph_format.space_after = Pt(2)
                current_paragraph = None
                continue
            
            # H3 headers (###)
            if line.startswith('### '):
                heading_text = line[4:].strip()
                heading = doc.add_heading(heading_text, level=3)
                heading_run = heading.runs[0]
                heading_run.font.size = Pt(10)  # Reduced from 15 to 10
                heading_run.font.color.rgb = RGBColor(167, 139, 250)
                heading.paragraph_format.space_before = Pt(3)
                heading.paragraph_format.space_after = Pt(2)
                current_paragraph = None
                continue
            
            # Horizontal rule (---)
            if line.startswith('---'):
                doc.add_paragraph('_' * 60)
                current_paragraph = None
                continue
            
            # Bullet points (-)
            if line.startswith('- '):
                bullet_text = line[2:].strip()
                para = doc.add_paragraph(style='List Bullet')
                self._add_styled_run(para, bullet_text)
                current_paragraph = None
                continue
            
            # Regular paragraph
            if not current_paragraph:
                current_paragraph = doc.add_paragraph()
            else:
                current_paragraph.add_run(' ')
            
            self._add_styled_run(current_paragraph, line)
    
    def _add_styled_run(self, paragraph, text):
        """Add text run with bold formatting (**text**)"""
        # Split by bold markers
        parts = re.split(r'(\*\*.*?\*\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                bold_text = part[2:-2]
                run = paragraph.add_run(bold_text)
                run.bold = True
            else:
                # Normal text
                paragraph.add_run(part)
    
    def _add_footer(self, doc, data):
        """Add footer to all pages"""
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Ensure UTF-8 string
        today = datetime.now().strftime('%d.%m.%Y')
        footer_text = f'AstroMind AI - Генерирано на {today} г. - Само за занимателна цел!'
        footer_run = footer_para.add_run(footer_text)
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(136, 136, 136)
    
    def _translate_sign(self, formatted_pos: str) -> str:
        """Translate English sign names to Bulgarian"""
        translated = formatted_pos
        for eng_sign, bg_sign in self.sign_names.items():
            translated = re.sub(eng_sign, bg_sign, translated, flags=re.IGNORECASE)
        return translated
    
    def _get_house_suffix(self, num: int) -> str:
        """Get Bulgarian house suffix"""
        if num == 1:
            return 'ви'
        elif num == 2:
            return 'ри'
        elif num == 3:
            return 'ти'
        elif num >= 4 and num <= 10:
            return 'ти'
        elif num == 11:
            return 'и'
        return 'ти'

